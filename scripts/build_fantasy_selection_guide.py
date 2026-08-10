from __future__ import annotations

import argparse
import itertools
import sqlite3
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


GUIDE_DATE = "10 августа 2026"

COLOR_BLUE = "2E74B5"
COLOR_BLUE_DARK = "1F4D78"
COLOR_BLUE_SOFT = "E8EEF5"
COLOR_BLUE_BOX = "F4F7FB"
COLOR_GRAY = "666666"
COLOR_BLACK = "111111"

ROLE_CONFIG = {
    "core_pair": {
        "label": "Пара core",
        "positions": [1, 3],
        "banner_scope": "core",
        "combo_template": {"red": 2, "green": 1},
    },
    "mid_single": {
        "label": "Mid",
        "positions": [2],
        "banner_scope": "mid",
        "combo_template": {"red": 1, "blue": 1, "green": 1},
    },
    "support_pair": {
        "label": "Пара support",
        "positions": [4, 5],
        "banner_scope": "support",
        "combo_template": {"blue": 2, "green": 1},
    },
}

COLOR_LABELS = {
    "red": "Красный",
    "blue": "Синий",
    "green": "Зеленый",
    "unknown": "Неизвестно",
}

STAT_DESCRIPTIONS = {
    "creep_score": "Фундаментальный фарм-стат для core и сильный общий ориентир по стабильности.",
    "gpm": "Надежная экономическая метрика, обычно одна из лучших красных опций для core.",
    "teamfight_participation": "Один из самых универсально полезных зеленых статов почти для любого слота.",
    "kills": "Больше помогает потолку, чем стабильности, поэтому редко должен быть единственным приоритетом.",
    "deaths": "В этой scoring-модели не так плох, как кажется интуитивно, из-за высокой базы.",
    "runes_grabbed": "Ключевой синий mid-стат, хорошо сочетается с красным фармом и teamfight.",
    "wards_placed": "Практически полезный support-стат, один из лучших синих ориентиров в текущем датасете.",
    "smokes_used": "Рабочий utility-стат, но интерпретировать его лучше аккуратно из-за особенностей источника.",
    "camps_stacked": "Неплохой синий стат для support-пары, особенно если важен хозяйственный вклад по карте.",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        child = tc_mar.find(qn(f"w:{key}"))
        if child is None:
            child = OxmlElement(f"w:{key}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_layout_fixed(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_run_font(run, *, size: int, bold=False, color=COLOR_BLACK, name="Calibri", italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, *, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    paragraph.alignment = align


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=4, line=1.0)
    r = p.add_run(title)
    set_run_font(r, size=24, bold=True, color=COLOR_BLUE_DARK)

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=10, line=1.15)
    r = p.add_run(subtitle)
    set_run_font(r, size=11, color=COLOR_GRAY)


def add_summary_box(doc: Document, bullets: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_layout_fixed(table)
    table.columns[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, COLOR_BLUE_BOX)
    set_cell_margins(cell, top=120, start=140, bottom=120, end=140)

    heading = cell.paragraphs[0]
    style_paragraph(heading, before=0, after=6, line=1.0)
    r = heading.add_run("Как читать этот гайд")
    set_run_font(r, size=12, bold=True, color=COLOR_BLUE_DARK)

    for bullet in bullets:
        p = cell.add_paragraph(style="List Bullet")
        style_paragraph(p, before=0, after=4, line=1.15)
        r = p.add_run(bullet)
        set_run_font(r, size=11)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, before=16, after=8, line=1.0)
        r = p.add_run(text)
        set_run_font(r, size=16, bold=True, color=COLOR_BLUE)
    elif level == 2:
        style_paragraph(p, before=12, after=6, line=1.0)
        r = p.add_run(text)
        set_run_font(r, size=13, bold=True, color=COLOR_BLUE)
    else:
        style_paragraph(p, before=8, after=4, line=1.0)
        r = p.add_run(text)
        set_run_font(r, size=12, bold=True, color=COLOR_BLUE_DARK)


def add_body_paragraph(doc: Document, text: str, *, muted=False) -> None:
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=6, line=1.2)
    r = p.add_run(text)
    set_run_font(r, size=11, color=COLOR_GRAY if muted else COLOR_BLACK)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, before=0, after=4, line=1.2)
        r = p.add_run(item)
        set_run_font(r, size=11)


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_layout_fixed(table)
    table.style = "Table Grid"

    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].width = Inches(widths[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[idx], COLOR_BLUE_SOFT)
        set_cell_margins(hdr[idx])
        p = hdr[idx].paragraphs[0]
        style_paragraph(p, before=0, after=0, line=1.0)
        r = p.add_run(header)
        set_run_font(r, size=10, bold=True, color=COLOR_BLUE_DARK)

    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].width = Inches(widths[idx])
            row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row[idx])
            p = row[idx].paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.1)
            r = p.add_run(str(value))
            set_run_font(r, size=10)


def discrete_p75(values: pd.Series) -> float:
    arr = sorted(float(v) for v in values.dropna().tolist())
    if not arr:
        return 0.0
    n = len(arr)
    i1 = int((n - 1) * 0.75)
    i2 = min(i1 + 1, n - 1)
    return (arr[i1] + arr[i2]) / 2.0


def format_stat_name(stat_name: str) -> str:
    return stat_name.replace("_", " ")


def fetch_banner_profile(con: sqlite3.Connection) -> dict[str, list[dict]]:
    df = pd.read_sql_query(
        """
        SELECT role_scope, banner_slot, stat_name, multiplier, notes
        FROM analytics_scoring_formula
        WHERE profile_id = 'my_current_banner_official_roles'
          AND enabled = 1
        ORDER BY role_scope, banner_slot
        """,
        con,
    )
    out: dict[str, list[dict]] = {}
    for _, row in df.iterrows():
        out.setdefault(row["role_scope"], []).append(row.to_dict())
    return out


def build_role_map_stat_df(con: sqlite3.Connection, role_key: str) -> pd.DataFrame:
    cfg = ROLE_CONFIG[role_key]
    pos_list = ", ".join(map(str, cfg["positions"]))
    expected_positions = len(cfg["positions"])
    sql = f"""
    WITH role_players AS (
        SELECT
            pir.account_id,
            pir.team_name,
            pir.official_name,
            pir.official_position
        FROM player_identity_registry pir
        WHERE pir.official_position IN ({pos_list})
          AND EXISTS (
              SELECT 1
              FROM analytics_ti2026_teams ti
              WHERE ti.team_name = pir.team_name
          )
    ),
    role_names AS (
        SELECT
            rp.team_name,
            GROUP_CONCAT(rp.official_name, ', ') AS player_names
        FROM (
            SELECT team_name, official_name, official_position
            FROM role_players
            ORDER BY team_name, official_position
        ) rp
        GROUP BY rp.team_name
        HAVING COUNT(DISTINCT rp.official_position) = {expected_positions}
    ),
    target_maps AS (
        SELECT
            f.match_id,
            f.team_name
        FROM player_game_fantasy_summary f
        JOIN role_players rp
          ON rp.account_id = f.account_id
         AND rp.team_name = f.team_name
        GROUP BY f.match_id, f.team_name
        HAVING COUNT(DISTINCT rp.official_position) = {expected_positions}
    )
    SELECT
        tm.match_id,
        tm.team_name,
        rn.player_names,
        sc.stat_name,
        COALESCE(sc.emblem_color, 'unknown') AS color_group,
        AVG(COALESCE(sp.base_points, 0.0)) AS points_x1
    FROM target_maps tm
    JOIN role_players rp
      ON rp.team_name = tm.team_name
    JOIN role_names rn
      ON rn.team_name = tm.team_name
    JOIN fantasy_scoring_stat_catalog sc
      ON 1 = 1
    LEFT JOIN fantasy_player_map_stat_points sp
      ON sp.match_id = tm.match_id
     AND sp.account_id = rp.account_id
     AND sp.team_name = rp.team_name
     AND sp.stat_name = sc.stat_name
    GROUP BY tm.match_id, tm.team_name, rn.player_names, sc.stat_name, sc.emblem_color
    """
    return pd.read_sql_query(sql, con)


def summarize_individual_stats(role_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for (stat_name, color_group), group in role_df.groupby(["stat_name", "color_group"], sort=False):
        rows.append(
            {
                "stat_name": stat_name,
                "color_group": color_group,
                "maps": int(group["match_id"].nunique()),
                "avg_x1": round(group["points_x1"].mean(), 2),
                "max_x1": round(group["points_x1"].max(), 2),
                "p75_x1": round(discrete_p75(group["points_x1"]), 2),
            }
        )
    out = pd.DataFrame(rows)
    return out.sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False).reset_index(drop=True)


def top_team_combos_for_single_stat(role_df: pd.DataFrame, stat_name: str) -> pd.DataFrame:
    df = role_df[role_df["stat_name"] == stat_name].copy()
    rows = []
    for (team_name, player_names), group in df.groupby(["team_name", "player_names"], sort=False):
        rows.append(
            {
                "team_name": team_name,
                "player_names": player_names,
                "maps": int(group["match_id"].nunique()),
                "avg_x1": round(group["points_x1"].mean(), 2),
                "max_x1": round(group["points_x1"].max(), 2),
                "p75_x1": round(discrete_p75(group["points_x1"]), 2),
            }
        )
    out = pd.DataFrame(rows)
    return out.sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False).reset_index(drop=True)


def generate_stat_combinations(summary_df: pd.DataFrame, role_key: str) -> list[tuple[str, ...]]:
    template = ROLE_CONFIG[role_key]["combo_template"]
    by_color: dict[str, list[str]] = {}
    positive = summary_df[summary_df["p75_x1"] > 0].copy()
    for color, take_n in template.items():
        stats = (
            positive[positive["color_group"] == color]
            .sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False)
            .head(5)["stat_name"]
            .tolist()
        )
        if len(stats) < take_n:
            return []
        by_color[color] = stats

    per_color_combos = [list(itertools.combinations(by_color[color], take_n)) for color, take_n in template.items()]
    all_combos = []
    for choice in itertools.product(*per_color_combos):
        flattened = tuple(sorted({stat for part in choice for stat in part}))
        if flattened not in all_combos:
            all_combos.append(flattened)
    return all_combos


def summarize_stat_combinations(role_df: pd.DataFrame, role_key: str) -> tuple[pd.DataFrame, dict[str, pd.DataFrame]]:
    combos = generate_stat_combinations(summarize_individual_stats(role_df), role_key)
    combo_summary_rows = []
    combo_team_tables: dict[str, pd.DataFrame] = {}

    for combo in combos:
        combo_df = role_df[role_df["stat_name"].isin(combo)].copy()
        if combo_df.empty:
            continue

        per_map = (
            combo_df.groupby(["match_id", "team_name", "player_names"], as_index=False)["points_x1"]
            .sum()
            .rename(columns={"points_x1": "combo_score"})
        )
        combo_name = " + ".join(format_stat_name(x) for x in combo)

        combo_summary_rows.append(
            {
                "combo_name": combo_name,
                "stats_count": len(combo),
                "avg_x1": round(per_map["combo_score"].mean(), 2),
                "max_x1": round(per_map["combo_score"].max(), 2),
                "p75_x1": round(discrete_p75(per_map["combo_score"]), 2),
            }
        )

        team_rows = []
        for (team_name, player_names), group in per_map.groupby(["team_name", "player_names"], sort=False):
            team_rows.append(
                {
                    "team_name": team_name,
                    "player_names": player_names,
                    "maps": int(group["match_id"].nunique()),
                    "avg_x1": round(group["combo_score"].mean(), 2),
                    "max_x1": round(group["combo_score"].max(), 2),
                    "p75_x1": round(discrete_p75(group["combo_score"]), 2),
                }
            )
        combo_team_tables[combo_name] = (
            pd.DataFrame(team_rows)
            .sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False)
            .reset_index(drop=True)
        )

    combo_summary_df = pd.DataFrame(combo_summary_rows)
    if combo_summary_df.empty:
        return combo_summary_df, combo_team_tables
    combo_summary_df = combo_summary_df.sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False).reset_index(drop=True)
    return combo_summary_df, combo_team_tables


def fetch_profile_rankings(con: sqlite3.Connection) -> dict[str, pd.DataFrame]:
    sql = """
    WITH base AS (
        SELECT
            team_name,
            core_players,
            mid_player,
            support_players,
            avg_core_fantasy_score AS core_score,
            mid_fantasy_score AS mid_score,
            avg_support_fantasy_score AS support_score
        FROM analytics_team_role_maps
        WHERE ti2026_qualified = 1
    ),
    unioned AS (
        SELECT 'core_pair' AS role_slot, team_name, core_players AS player_names, core_score AS score FROM base
        UNION ALL
        SELECT 'mid_single', team_name, mid_player AS player_names, mid_score AS score FROM base
        UNION ALL
        SELECT 'support_pair', team_name, support_players AS player_names, support_score AS score FROM base
    ),
    ranked AS (
        SELECT
            role_slot,
            team_name,
            player_names,
            score,
            ROW_NUMBER() OVER (PARTITION BY role_slot, team_name ORDER BY score) AS rn,
            COUNT(*) OVER (PARTITION BY role_slot, team_name) AS cnt
        FROM unioned
    )
    SELECT
        role_slot,
        team_name,
        player_names,
        COUNT(*) AS maps_played,
        ROUND(AVG(score), 2) AS avg_score,
        ROUND(MAX(score), 2) AS max_score,
        ROUND(
            AVG(
                CASE
                    WHEN rn IN (
                        CAST(((cnt - 1) * 0.75) AS INTEGER) + 1,
                        CAST(((cnt - 1) * 0.75) AS INTEGER) + 2
                    )
                    THEN score
                END
            ),
            2
        ) AS p75_score
    FROM ranked
    GROUP BY role_slot, team_name, player_names
    ORDER BY role_slot, p75_score DESC, max_score DESC
    """
    df = pd.read_sql_query(sql, con)
    return {
        role: grp.reset_index(drop=True)
        for role, grp in df.groupby("role_slot", sort=False)
    }


def fetch_support_caveat(con: sqlite3.Connection) -> str:
    row = con.execute(
        """
        SELECT caveat
        FROM analytics_support_caveat
        WHERE role_key = 'support_pair'
        LIMIT 1
        """
    ).fetchone()
    return row[0] if row else ""


def build_document(db_path: Path, out_path: Path) -> None:
    con = sqlite3.connect(str(db_path))
    banner_profile = fetch_banner_profile(con)
    profile_rankings = fetch_profile_rankings(con)
    support_caveat = fetch_support_caveat(con)
    qualified_teams = con.execute("SELECT COUNT(*) FROM analytics_ti2026_teams").fetchone()[0]

    role_payload = {}
    for role_key in ROLE_CONFIG:
        role_df = build_role_map_stat_df(con, role_key)
        stat_summary = summarize_individual_stats(role_df)
        combo_summary, combo_team_tables = summarize_stat_combinations(role_df, role_key)
        top_stat_team_tables = {
            row["stat_name"]: top_team_combos_for_single_stat(role_df, row["stat_name"])
            for _, row in stat_summary.head(3).iterrows()
        }
        role_payload[role_key] = {
            "role_df": role_df,
            "stat_summary": stat_summary,
            "combo_summary": combo_summary,
            "combo_team_tables": combo_team_tables,
            "top_stat_team_tables": top_stat_team_tables,
        }
    con.close()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    footer = section.footer.paragraphs[0]
    style_paragraph(footer, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    r = footer.add_run("Гайд по выбору фэнтези на EWC 2026")
    set_run_font(r, size=9, color=COLOR_GRAY)

    add_title(
        doc,
        "Гайд по выбору фэнтези на EWC 2026",
        f"Человеко-ориентированный обзор на основе базы проекта. Охват: {qualified_teams} TI 2026 qualified teams. Актуальность: {GUIDE_DATE}.",
    )

    add_summary_box(
        doc,
        [
            "Сначала идет общее введение, потом аналитика по статам x1.0, затем лучшие stat-комбинации x1.0 и уже после этого рекомендации по слот-комбинациям игроков.",
            "Главный ориентир для shortlisting - p75: он лучше average показывает, кто стабильно добирается до верхней части распределения.",
            "Почти вся основная часть разделена по ролям, чтобы документ не превращался в одну перегруженную сводку.",
            "Текущий баннерный профиль вынесен в конец как пример практического применения, а не как главный способ читать весь гайд.",
        ],
    )

    add_heading(doc, "1. Введение", 1)
    add_bullets(
        doc,
        [
            "Этот гайд описывает не просто игроков, а слоты выбора: core_pair, mid_single и support_pair.",
            "Внутри слота сначала полезно понять, какие отдельные статы сильнее, потом какие stat-комбинации работают лучше, и только затем уже выбирать конкретные команды и пары игроков.",
            "Для core и mid рекомендации заметно надежнее. Для support интерпретацию лучше держать осторожнее.",
            support_caveat or "Support-слой в этой базе менее надежен, чем core и mid.",
        ],
    )

    for role_key in ["core_pair", "mid_single", "support_pair"]:
        payload = role_payload[role_key]
        role_label = ROLE_CONFIG[role_key]["label"]
        add_heading(doc, f"2. {role_label}", 1)

        add_heading(doc, "2.1 Ранжировка отдельных статистик (x1.0)", 2)
        add_body_paragraph(
            doc,
            "Сначала смотри на одиночные статы внутри роли. Здесь полезнее всего p75, потому что он лучше отделяет действительно рабочие направления от красивых, но редких всплесков.",
        )
        stat_rows = []
        for _, row in payload["stat_summary"].head(8).iterrows():
            stat_rows.append(
                [
                    format_stat_name(row["stat_name"]),
                    COLOR_LABELS.get(row["color_group"], row["color_group"]),
                    f'{row["p75_x1"]:.2f}',
                    f'{row["avg_x1"]:.2f}',
                    f'{row["max_x1"]:.2f}',
                    STAT_DESCRIPTIONS.get(row["stat_name"], "Полезный ситуативный стат в текущей выборке."),
                ]
            )
        add_simple_table(
            doc,
            ["Стат", "Цвет", "P75 x1", "Avg x1", "Max x1", "Комментарий"],
            stat_rows,
            widths=[1.3, 0.75, 0.75, 0.75, 0.75, 2.2],
        )

        add_heading(doc, "2.2 Лучшие сочетания статистик (x1.0)", 2)
        add_body_paragraph(
            doc,
            "Дальше смотрим уже не на отдельные статы, а на их сочетания. Здесь все статы считаются с коэффициентами 1.0, чтобы увидеть именно силу комбинации как таковой.",
        )
        combo_rows = []
        for _, row in payload["combo_summary"].head(6).iterrows():
            combo_rows.append(
                [
                    row["combo_name"],
                    str(int(row["stats_count"])),
                    f'{row["p75_x1"]:.2f}',
                    f'{row["avg_x1"]:.2f}',
                    f'{row["max_x1"]:.2f}',
                ]
            )
        add_simple_table(
            doc,
            ["Сочетание статов", "Статов", "P75 x1", "Avg x1", "Max x1"],
            combo_rows,
            widths=[3.1, 0.6, 0.8, 0.8, 0.8],
        )

        add_heading(doc, "2.3 Комбинации игроков для топовых отдельных статов", 2)
        add_body_paragraph(
            doc,
            "Ниже показаны сочетания игроков, которые особенно хорошо конвертируют самые сильные отдельные статы этой роли.",
        )
        for stat_name, top_df in payload["top_stat_team_tables"].items():
            add_heading(doc, format_stat_name(stat_name), 3)
            rows = []
            for _, row in top_df.head(5).iterrows():
                rows.append(
                    [
                        row["team_name"],
                        row["player_names"],
                        f'{row["p75_x1"]:.2f}',
                        f'{row["avg_x1"]:.2f}',
                        f'{row["max_x1"]:.2f}',
                    ]
                )
            add_simple_table(
                doc,
                ["Команда", "Игроки", "P75 x1", "Avg x1", "Max x1"],
                rows,
                widths=[1.45, 2.1, 0.85, 0.85, 0.85],
            )

        add_heading(doc, "2.4 Комбинации игроков для лучших сочетаний статов", 2)
        add_body_paragraph(
            doc,
            "Теперь тот же вопрос, но уже для лучших stat-комбинаций x1.0. Это полезнее, когда ты реально выбираешь направление баннера, а не один конкретный стат.",
        )
        top_combo_names = payload["combo_summary"].head(3)["combo_name"].tolist()
        for combo_name in top_combo_names:
            add_heading(doc, combo_name, 3)
            top_df = payload["combo_team_tables"][combo_name]
            rows = []
            for _, row in top_df.head(5).iterrows():
                rows.append(
                    [
                        row["team_name"],
                        row["player_names"],
                        f'{row["p75_x1"]:.2f}',
                        f'{row["avg_x1"]:.2f}',
                        f'{row["max_x1"]:.2f}',
                        str(int(row["maps"])),
                    ]
                )
            add_simple_table(
                doc,
                ["Команда", "Игроки", "P75 x1", "Avg x1", "Max x1", "Карт"],
                rows,
                widths=[1.35, 1.85, 0.8, 0.8, 0.8, 0.6],
            )

    add_heading(doc, "3. Пример: текущий баннерный профиль", 1)
    add_body_paragraph(
        doc,
        "Эта секция уже не про абстрактные x1.0-сочетания, а про твой конкретный сохраненный баннерный профиль. Ее лучше читать как пример практического применения всей логики выше.",
    )

    for scope in ["core", "mid", "support"]:
        rows = []
        for item in banner_profile.get(scope, []):
            rows.append(
                [
                    str(item["banner_slot"]),
                    format_stat_name(item["stat_name"]),
                    f'x{float(item["multiplier"]):.1f}',
                    item["notes"] or "",
                ]
            )
        add_heading(doc, f"Профиль {scope}", 2)
        add_simple_table(
            doc,
            ["Слот", "Стат", "Множитель", "Комментарий"],
            rows,
            widths=[0.7, 1.8, 0.8, 3.2],
        )

    add_heading(doc, "4. Рекомендации по текущему баннеру", 1)
    add_body_paragraph(
        doc,
        "Ниже показаны slot-комбинации, которые лучше всего выглядят именно под текущий профиль. Это уже не заменяет общую аналитику выше, а накладывается на нее.",
    )
    for role_key in ["core_pair", "mid_single", "support_pair"]:
        role_label = ROLE_CONFIG[role_key]["label"]
        add_heading(doc, role_label, 2)
        rows = []
        df = profile_rankings.get(role_key, pd.DataFrame()).head(6)
        for _, row in df.iterrows():
            rows.append(
                [
                    row["team_name"],
                    row["player_names"],
                    f'{row["p75_score"]:.2f}',
                    f'{row["avg_score"]:.2f}',
                    f'{row["max_score"]:.2f}',
                    str(int(row["maps_played"])),
                ]
            )
        add_simple_table(
            doc,
            ["Команда", "Игроки", "P75", "Avg", "Max", "Карт"],
            rows,
            widths=[1.35, 1.85, 0.8, 0.8, 0.8, 0.6],
        )

    add_heading(doc, "5. Как пользоваться этим документом", 1)
    add_bullets(
        doc,
        [
            "Если хочешь понять общий приоритет роллов - смотри сначала на одиночные stat rankings внутри роли.",
            "Если хочешь понять лучшие направления баннера как набора - смотри на stat-комбинации x1.0.",
            "Если хочешь быстро выбрать реальный слот команды - переходи к блокам с player combinations под отдельные статы и stat-комбинации.",
            "Если хочешь увидеть, что говорит именно твой текущий профиль - открывай самый конец документа.",
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--db", required=True, type=Path)
    parser.add_argument("--out", required=True, type=Path)
    args = parser.parse_args()
    build_document(args.db, args.out)
    print(args.out)


if __name__ == "__main__":
    main()

from __future__ import annotations

import argparse
import itertools
import sqlite3
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT / "src") not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT / "src"))

from project_db import canonical_db_path, resolve_db_path  # noqa: E402

DEFAULT_DB_PATH = canonical_db_path(PROJECT_ROOT)
DEFAULT_DOCX_PATH = PROJECT_ROOT / "docs" / "EWC2026_Fantasy_Selection_Guide.docx"
DEFAULT_MD_PATH = PROJECT_ROOT / "docs" / "EWC2026_Fantasy_Selection_Guide.md"

ROLE_CONFIG = {
    "core_pair": {
        "label": "Core Pair",
        "positions": [1, 3],
        "banner_scope": "core",
        "combo_template": {"red": 2, "green": 1},
    },
    "mid_single": {
        "label": "Mid",
        "positions": [2],
        "banner_scope": "mid",
        "combo_template": {"red": 1, "blue": 1, "green": 1},
    },
    "support_pair": {
        "label": "Support Pair",
        "positions": [4, 5],
        "banner_scope": "support",
        "combo_template": {"blue": 2, "green": 1},
    },
}

ROLE_CATEGORY_BY_KEY = {
    "core_pair": "core_avg",
    "mid_single": "mid",
    "support_pair": "support_avg",
}

COLOR_LABELS = {
    "red": "Red",
    "blue": "Blue",
    "green": "Green",
    "unknown": "Unknown",
}

SUPPORTED_COVERAGE_STATUSES = ("filled_backfill", "filled_existing", "filled_approximation")

POINTS_COLUMN_TO_STAT = {
    "kills_points": "kills",
    "deaths_points": "deaths",
    "creep_score_points": "creep_score",
    "gpm_points": "gpm",
    "wards_points": "wards_placed",
    "camps_stacked_points": "camps_stacked",
    "runes_grabbed_points": "runes_grabbed",
    "watchers_taken_points": "watchers_taken",
    "lotus_points": "lotus",
    "roshan_points": "roshan_kills",
    "teamfight_participation_points": "teamfight_participation",
    "stuns_points": "stuns",
    "tormentor_points": "tormentor_kills",
    "courier_points": "courier_kills",
    "first_blood_points": "first_blood",
    "smokes_points": "smokes_used",
}

TEAM_ABBREVIATIONS = {
    "1w": "1W",
    "_PowerRangers": "PR",
    "Aurora Gaming": "AUR",
    "BetBoom Team": "BB",
    "BoomBoys": "BBY",
    "Gaimin Gladiators": "GG",
    "GamerLegion": "GL",
    "LGD Gaming": "LGD",
    "Nigma Galaxy": "NGX",
    "OG": "OG",
    "PVISION": "PV",
    "Rune Eaters": "RE",
    "Team Falcons": "FLC",
    "Team Liquid": "TL",
    "Team Spirit": "TS",
    "Team Yandex": "TY",
    "Tundra Esports": "TND",
    "Vici Gaming": "VG",
    "Virtus.pro": "VP",
    "Xtreme Gaming": "XG",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a clean fantasy selection guide from the compact SQLite database.")
    parser.add_argument("--db-path", default="", help="Path to the compact SQLite database.")
    parser.add_argument("--docx-out", default=str(DEFAULT_DOCX_PATH), help="Output DOCX path.")
    parser.add_argument("--md-out", default=str(DEFAULT_MD_PATH), help="Output Markdown path.")
    return parser.parse_args()


def make_grid_table(doc: Document, rows: int, cols: int):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
    return table


def discrete_p75(values: pd.Series) -> float:
    arr = sorted(float(v) for v in values.dropna().tolist())
    if not arr:
        return 0.0
    n = len(arr)
    i1 = int((n - 1) * 0.75)
    i2 = min(i1 + 1, n - 1)
    return (arr[i1] + arr[i2]) / 2.0


def stat_data_confidence(row: pd.Series) -> int:
    coverage_status = row.get("coverage_status", "")
    stat_name = row.get("stat_name", "")
    preferred_source = row.get("preferred_source", "")

    if coverage_status == "filled_existing":
        return 10
    if coverage_status == "filled_backfill":
        if preferred_source == "opendota":
            return 9
        if preferred_source == "source2_demo":
            return 8
        return 8
    if coverage_status == "filled_approximation":
        if stat_name == "tormentor_kills":
            return 6
        return 7
    if coverage_status == "source_needed":
        return 2
    return 5


def format_stat_name(stat_name: str) -> str:
    return stat_name.replace("_", " ")


def abbreviate_team_name(team_name: str) -> str:
    if team_name in TEAM_ABBREVIATIONS:
        return TEAM_ABBREVIATIONS[team_name]
    tokens = [token for token in team_name.replace(".", " ").replace("_", " ").split() if token]
    if not tokens:
        return team_name
    if len(tokens) == 1:
        return tokens[0][:4].upper()
    return "".join(token[0].upper() for token in tokens[:4])


def build_role_map_stat_df(con: sqlite3.Connection, role_key: str) -> pd.DataFrame:
    role_category = ROLE_CATEGORY_BY_KEY[role_key]
    wide = pd.read_sql_query(
        """
        SELECT rms.*, ti.team_name AS ti_team_name
        FROM player_map_role_category_stats rms
        JOIN analytics_ti2026_teams ti
          ON ti.team_name = rms.team_name
        WHERE rms.role_category = ?
        """,
        con,
        params=[role_category],
    )
    catalog = pd.read_sql_query(
        """
        SELECT
            sc.stat_name,
            COALESCE(sc.emblem_color, 'unknown') AS emblem_color,
            COALESCE(cov.coverage_status, 'filled_existing') AS coverage_status,
            COALESCE(cov.preferred_source, 'sqlite') AS preferred_source
        FROM fantasy_scoring_stat_catalog sc
        LEFT JOIN analytics_fantasy_backfill_coverage cov
          ON cov.stat_name = sc.stat_name
        WHERE COALESCE(cov.coverage_status, 'filled_existing') IN ('filled_backfill', 'filled_existing', 'filled_approximation')
        """,
        con,
    )
    catalog_map = {
        row["stat_name"]: {
            "color_group": row["emblem_color"],
            "coverage_status": row["coverage_status"],
            "preferred_source": row["preferred_source"],
        }
        for _, row in catalog.iterrows()
    }

    rows: list[dict[str, object]] = []
    for _, row in wide.iterrows():
        for points_col, stat_name in POINTS_COLUMN_TO_STAT.items():
            meta = catalog_map.get(stat_name)
            if meta is None:
                continue
            rows.append(
                {
                    "match_id": int(row["match_id"]),
                    "team_name": row["team_name"],
                    "player_names": row["player_names"],
                    "stat_name": stat_name,
                    "color_group": meta["color_group"],
                    "coverage_status": meta["coverage_status"],
                    "preferred_source": meta["preferred_source"],
                    "points_x1": float(row[points_col] or 0.0),
                }
            )
    out = pd.DataFrame(rows)
    positive_stats = set(out.groupby("stat_name")["points_x1"].max().loc[lambda s: s > 0].index.tolist())
    return out[out["stat_name"].isin(positive_stats)].reset_index(drop=True)


def summarize_individual_stats(role_df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for (stat_name, color_group, coverage_status, preferred_source), group in role_df.groupby(
        ["stat_name", "color_group", "coverage_status", "preferred_source"], sort=False
    ):
        rows.append(
            {
                "stat_name": stat_name,
                "color_group": color_group,
                "coverage_status": coverage_status,
                "preferred_source": preferred_source,
                "maps": int(group["match_id"].nunique()),
                "avg_x1": round(group["points_x1"].mean(), 2),
                "max_x1": round(group["points_x1"].max(), 2),
                "p75_x1": round(discrete_p75(group["points_x1"]), 2),
                "active_rate": round(float((group["points_x1"] > 0).mean()), 4),
            }
        )
    out = pd.DataFrame(rows)
    return out.sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False).reset_index(drop=True)


def top_team_combos_for_single_stat(role_df: pd.DataFrame, stat_name: str) -> pd.DataFrame:
    df = role_df[role_df["stat_name"] == stat_name].copy()
    rows = []
    for (team_name, player_names), group in df.groupby(["team_name", "player_names"], sort=False):
        rows.append(
            {
                "team_name": team_name,
                "player_names": player_names,
                "maps": int(group["match_id"].nunique()),
                "avg_x1": round(group["points_x1"].mean(), 2),
                "max_x1": round(group["points_x1"].max(), 2),
                "p75_x1": round(discrete_p75(group["points_x1"]), 2),
            }
        )
    if not rows:
        return pd.DataFrame(columns=["team_name", "player_names", "maps", "avg_x1", "max_x1", "p75_x1"])
    return pd.DataFrame(rows).sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False).reset_index(drop=True)


def compact_best_of_class(top_df: pd.DataFrame, limit: int = 3) -> str:
    if top_df.empty:
        return "-"
    parts: list[str] = []
    for _, row in top_df.head(limit).iterrows():
        parts.append(f"{row['player_names']} [{abbreviate_team_name(row['team_name'])}]")
    return "\n".join(parts)


def generate_stat_combinations(summary_df: pd.DataFrame, role_key: str) -> list[tuple[str, ...]]:
    template = ROLE_CONFIG[role_key]["combo_template"]
    positive = summary_df[summary_df["p75_x1"] > 0].copy()
    by_color: dict[str, list[str]] = {}
    for color, take_n in template.items():
        stats = (
            positive[positive["color_group"] == color]
            .sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False)
            .head(5)["stat_name"]
            .tolist()
        )
        if len(stats) < take_n:
            return []
        by_color[color] = stats
    per_color = [list(itertools.combinations(by_color[color], take_n)) for color, take_n in template.items()]
    combos: list[tuple[str, ...]] = []
    for choice in itertools.product(*per_color):
        flattened = tuple(sorted({stat for part in choice for stat in part}))
        if flattened not in combos:
            combos.append(flattened)
    return combos


def summarize_stat_combinations(role_df: pd.DataFrame, role_key: str) -> tuple[pd.DataFrame, dict[str, pd.DataFrame]]:
    combos = generate_stat_combinations(summarize_individual_stats(role_df), role_key)
    combo_rows = []
    combo_team_tables: dict[str, pd.DataFrame] = {}
    for combo in combos:
        combo_df = role_df[role_df["stat_name"].isin(combo)].copy()
        if combo_df.empty:
            continue
        per_map = (
            combo_df.groupby(["match_id", "team_name", "player_names"], as_index=False)["points_x1"]
            .sum()
            .rename(columns={"points_x1": "combo_score"})
        )
        combo_name = " + ".join(format_stat_name(x) for x in combo)
        combo_rows.append(
            {
                "combo_name": combo_name,
                "stats_count": len(combo),
                "avg_x1": round(per_map["combo_score"].mean(), 2),
                "max_x1": round(per_map["combo_score"].max(), 2),
                "p75_x1": round(discrete_p75(per_map["combo_score"]), 2),
            }
        )
        team_rows = []
        for (team_name, player_names), group in per_map.groupby(["team_name", "player_names"], sort=False):
            team_rows.append(
                {
                    "team_name": team_name,
                    "player_names": player_names,
                    "maps": int(group["match_id"].nunique()),
                    "avg_x1": round(group["combo_score"].mean(), 2),
                    "max_x1": round(group["combo_score"].max(), 2),
                    "p75_x1": round(discrete_p75(group["combo_score"]), 2),
                }
            )
        combo_team_tables[combo_name] = pd.DataFrame(team_rows).sort_values(
            ["p75_x1", "max_x1", "avg_x1"], ascending=False
        ).reset_index(drop=True)
    if not combo_rows:
        return pd.DataFrame(columns=["combo_name", "stats_count", "avg_x1", "max_x1", "p75_x1"]), combo_team_tables
    combo_df = pd.DataFrame(combo_rows).sort_values(["p75_x1", "max_x1", "avg_x1"], ascending=False).reset_index(drop=True)
    return combo_df, combo_team_tables


def fetch_banner_profile(con: sqlite3.Connection) -> pd.DataFrame:
    return pd.read_sql_query(
        """
        SELECT profile_id, role_scope, banner_slot, stat_name, multiplier, notes
        FROM analytics_scoring_formula
        WHERE profile_id = 'my_current_banner_official_roles'
          AND enabled = 1
        ORDER BY role_scope, banner_slot
        """,
        con,
    )


def fetch_optimizer_recommendations(con: sqlite3.Connection) -> pd.DataFrame:
    return pd.read_sql_query(
        """
        SELECT role_slot, team_name, player_names, optimizer_score_1_100,
               predicted_score_raw, best2_series_score, p75_series_score
        FROM analytics_optimizer_role_slots
        WHERE run_id = 'optimizer_my_current_banner_official_roles_all_default'
          AND profile_id = 'my_current_banner_official_roles'
          AND ti2026_qualified = 1
        ORDER BY role_slot, optimizer_score_1_100 DESC, predicted_score_raw DESC
        """,
        con,
    )


def write_markdown(
    out_path: Path,
    *,
    qualified_teams: int,
    player_maps: int,
    role_payload: dict[str, dict[str, object]],
    banner_profile: pd.DataFrame,
    optimizer_df: pd.DataFrame,
) -> None:
    lines: list[str] = []
    lines.append("# EWC 2026 Fantasy Selection Guide")
    lines.append("")
    lines.append(
        f"Guide snapshot: `{qualified_teams}` TI 2026 qualified teams, `{player_maps}` player-map rows, compact database updated on `2026-08-11`."
    )
    lines.append("")
    lines.append("## How to read this guide")
    lines.append("")
    lines.append("- Start with the single-stat tables inside each role. For fantasy shortlists, `p75` is usually more informative than the mean.")
    lines.append("- Then look at stat combinations with all multipliers set to `1.0` to understand which bundles of stats produce the strongest repeatable value.")
    lines.append("- Only after that should you decide which exact team/player combinations fit your current banner profile.")
    lines.append("- Support stats are included in the default analysis too; just read utility-heavy support metrics a bit more contextually than core farm stats.")
    lines.append("")
    for role_key in ["core_pair", "mid_single", "support_pair"]:
        cfg = ROLE_CONFIG[role_key]
        stat_summary = role_payload[role_key]["stat_summary"]
        combo_summary = role_payload[role_key]["combo_summary"]
        combo_team_tables = role_payload[role_key]["combo_team_tables"]
        stat_team_tables = role_payload[role_key]["all_stat_team_tables"]

        lines.append(f"## {cfg['label']}")
        lines.append("")
        lines.append("### 2.1 Single-stat ranking (`x1.0`)")
        lines.append("")
        lines.append("| Stat | Color | P75 x1 | Avg x1 | Max x1 | Trust | Best of class |")
        lines.append("|---|---:|---:|---:|---:|---:|---|")
        for _, row in stat_summary.iterrows():
            lines.append(
                f"| {format_stat_name(row['stat_name'])} | {COLOR_LABELS.get(row['color_group'], row['color_group'])} | "
                f"{row['p75_x1']:.2f} | {row['avg_x1']:.2f} | {row['max_x1']:.2f} | "
                f"{stat_data_confidence(row)} | {compact_best_of_class(stat_team_tables[row['stat_name']])} |"
            )
        lines.append("")
        lines.append("### 2.2 Best stat combinations (`x1.0`)")
        lines.append("")
        lines.append("| Stat combination | Stats | P75 x1 | Avg x1 | Max x1 |")
        lines.append("|---|---:|---:|---:|---:|")
        for _, row in combo_summary.head(8).iterrows():
            lines.append(
                f"| {row['combo_name']} | {int(row['stats_count'])} | {row['p75_x1']:.2f} | {row['avg_x1']:.2f} | {row['max_x1']:.2f} |"
            )
        lines.append("")
        lines.append("### 2.3 Best team/player combinations for featured single stats")
        lines.append("")
        for stat_name in stat_summary.head(3)["stat_name"].tolist():
            lines.append(f"#### {format_stat_name(stat_name)}")
            lines.append("")
            lines.append("| Team | Players | P75 x1 | Avg x1 | Max x1 |")
            lines.append("|---|---|---:|---:|---:|")
            for _, row in stat_team_tables[stat_name].head(5).iterrows():
                lines.append(
                    f"| {row['team_name']} | {row['player_names']} | {row['p75_x1']:.2f} | {row['avg_x1']:.2f} | {row['max_x1']:.2f} |"
                )
            lines.append("")
        lines.append("### 2.4 Best team/player combinations for featured stat-combos")
        lines.append("")
        for combo_name in combo_summary.head(3)["combo_name"].tolist():
            lines.append(f"#### {combo_name}")
            lines.append("")
            lines.append("| Team | Players | P75 x1 | Avg x1 | Max x1 |")
            lines.append("|---|---|---:|---:|---:|")
            for _, row in combo_team_tables[combo_name].head(5).iterrows():
                lines.append(
                    f"| {row['team_name']} | {row['player_names']} | {row['p75_x1']:.2f} | {row['avg_x1']:.2f} | {row['max_x1']:.2f} |"
                )
            lines.append("")

    lines.append("## Current banner profile example")
    lines.append("")
    lines.append("### Stored profile formula")
    lines.append("")
    lines.append("| Role | Slot | Stat | Multiplier | Notes |")
    lines.append("|---|---:|---|---:|---|")
    for _, row in banner_profile.iterrows():
        lines.append(
            f"| {row['role_scope']} | {int(row['banner_slot'])} | {row['stat_name']} | {float(row['multiplier']):.2f} | {row['notes'] or ''} |"
        )
    lines.append("")
    lines.append("### Optimizer suggestions for that profile")
    lines.append("")
    lines.append("| Role slot | Team | Players | Optimizer 1-100 | Predicted raw | Best2 series | P75 series |")
    lines.append("|---|---|---|---:|---:|---:|---:|")
    for _, row in optimizer_df.groupby("role_slot", sort=False).head(8).iterrows():
        lines.append(
            f"| {row['role_slot']} | {row['team_name']} | {row['player_names']} | "
            f"{row['optimizer_score_1_100']:.2f} | {row['predicted_score_raw']:.2f} | "
            f"{row['best2_series_score']:.2f} | {row['p75_series_score']:.2f} |"
        )
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_docx(
    out_path: Path,
    *,
    qualified_teams: int,
    player_maps: int,
    role_payload: dict[str, dict[str, object]],
    banner_profile: pd.DataFrame,
    optimizer_df: pd.DataFrame,
) -> None:
    doc = Document()
    doc.add_heading("EWC 2026 Fantasy Selection Guide", level=0)
    doc.add_paragraph(
        f"Guide snapshot: {qualified_teams} TI 2026 qualified teams, "
        f"{player_maps} player-map rows, compact database updated on 2026-08-11."
    )
    doc.add_heading("How to read this guide", level=1)
    for bullet in [
        "Start with the single-stat tables. For fantasy shortlists, p75 is usually more informative than the mean.",
        "Then look at stat combinations with all multipliers fixed at 1.0.",
        "Only after that should you choose exact teams and player combinations for your current profile.",
        "Support stats are part of the default analysis as well; utility-heavy support metrics simply deserve a bit more context when interpreted.",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    for role_key in ["core_pair", "mid_single", "support_pair"]:
        cfg = ROLE_CONFIG[role_key]
        stat_summary = role_payload[role_key]["stat_summary"]
        combo_summary = role_payload[role_key]["combo_summary"]
        combo_team_tables = role_payload[role_key]["combo_team_tables"]
        stat_team_tables = role_payload[role_key]["all_stat_team_tables"]

        doc.add_heading(cfg["label"], level=1)
        doc.add_heading("2.1 Single-stat ranking (x1.0)", level=2)
        table = make_grid_table(doc, rows=1, cols=7)
        headers = ["Stat", "Color", "P75 x1", "Avg x1", "Max x1", "Trust", "Best of class"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
        for _, row in stat_summary.iterrows():
            cells = table.add_row().cells
            cells[0].text = format_stat_name(row["stat_name"])
            cells[1].text = COLOR_LABELS.get(row["color_group"], row["color_group"])
            cells[2].text = f"{row['p75_x1']:.2f}"
            cells[3].text = f"{row['avg_x1']:.2f}"
            cells[4].text = f"{row['max_x1']:.2f}"
            cells[5].text = str(stat_data_confidence(row))
            cells[6].text = compact_best_of_class(stat_team_tables[row["stat_name"]])

        doc.add_heading("2.2 Best stat combinations (x1.0)", level=2)
        combo_table = make_grid_table(doc, rows=1, cols=5)
        for idx, header in enumerate(["Stat combination", "Stats", "P75 x1", "Avg x1", "Max x1"]):
            combo_table.rows[0].cells[idx].text = header
        for _, row in combo_summary.head(8).iterrows():
            cells = combo_table.add_row().cells
            cells[0].text = row["combo_name"]
            cells[1].text = str(int(row["stats_count"]))
            cells[2].text = f"{row['p75_x1']:.2f}"
            cells[3].text = f"{row['avg_x1']:.2f}"
            cells[4].text = f"{row['max_x1']:.2f}"

        doc.add_heading("2.3 Best team/player combinations for featured single stats", level=2)
        for stat_name in stat_summary.head(3)["stat_name"].tolist():
            doc.add_heading(format_stat_name(stat_name), level=3)
            stat_table = make_grid_table(doc, rows=1, cols=5)
            for idx, header in enumerate(["Team", "Players", "P75 x1", "Avg x1", "Max x1"]):
                stat_table.rows[0].cells[idx].text = header
            for _, row in stat_team_tables[stat_name].head(5).iterrows():
                cells = stat_table.add_row().cells
                cells[0].text = row["team_name"]
                cells[1].text = row["player_names"]
                cells[2].text = f"{row['p75_x1']:.2f}"
                cells[3].text = f"{row['avg_x1']:.2f}"
                cells[4].text = f"{row['max_x1']:.2f}"

        doc.add_heading("2.4 Best team/player combinations for featured stat-combos", level=2)
        for combo_name in combo_summary.head(3)["combo_name"].tolist():
            doc.add_heading(combo_name, level=3)
            combo_team_table = make_grid_table(doc, rows=1, cols=5)
            for idx, header in enumerate(["Team", "Players", "P75 x1", "Avg x1", "Max x1"]):
                combo_team_table.rows[0].cells[idx].text = header
            for _, row in combo_team_tables[combo_name].head(5).iterrows():
                cells = combo_team_table.add_row().cells
                cells[0].text = row["team_name"]
                cells[1].text = row["player_names"]
                cells[2].text = f"{row['p75_x1']:.2f}"
                cells[3].text = f"{row['avg_x1']:.2f}"
                cells[4].text = f"{row['max_x1']:.2f}"

    doc.add_heading("Current banner profile example", level=1)
    profile_table = make_grid_table(doc, rows=1, cols=5)
    for idx, header in enumerate(["Role", "Slot", "Stat", "Multiplier", "Notes"]):
        profile_table.rows[0].cells[idx].text = header
    for _, row in banner_profile.iterrows():
        cells = profile_table.add_row().cells
        cells[0].text = row["role_scope"]
        cells[1].text = str(int(row["banner_slot"]))
        cells[2].text = row["stat_name"]
        cells[3].text = f"{float(row['multiplier']):.2f}"
        cells[4].text = row["notes"] or ""

    doc.add_heading("Optimizer suggestions for that profile", level=2)
    opt_table = make_grid_table(doc, rows=1, cols=7)
    for idx, header in enumerate(["Role slot", "Team", "Players", "Optimizer 1-100", "Predicted raw", "Best2 series", "P75 series"]):
        opt_table.rows[0].cells[idx].text = header
    for _, row in optimizer_df.groupby("role_slot", sort=False).head(8).iterrows():
        cells = opt_table.add_row().cells
        cells[0].text = row["role_slot"]
        cells[1].text = row["team_name"]
        cells[2].text = row["player_names"]
        cells[3].text = f"{row['optimizer_score_1_100']:.2f}"
        cells[4].text = f"{row['predicted_score_raw']:.2f}"
        cells[5].text = f"{row['best2_series_score']:.2f}"
        cells[6].text = f"{row['p75_series_score']:.2f}"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    args = parse_args()
    db_path = resolve_db_path(PROJECT_ROOT, args.db_path or None).resolve()
    docx_out = Path(args.docx_out).resolve()
    md_out = Path(args.md_out).resolve()

    con = sqlite3.connect(str(db_path))
    role_payload: dict[str, dict[str, object]] = {}
    for role_key in ROLE_CONFIG:
        role_df = build_role_map_stat_df(con, role_key)
        role_payload[role_key] = {
            "role_df": role_df,
            "stat_summary": summarize_individual_stats(role_df),
        }
        combo_summary, combo_team_tables = summarize_stat_combinations(role_df, role_key)
        role_payload[role_key]["combo_summary"] = combo_summary
        role_payload[role_key]["combo_team_tables"] = combo_team_tables
        role_payload[role_key]["all_stat_team_tables"] = {
            row["stat_name"]: top_team_combos_for_single_stat(role_df, row["stat_name"])
            for _, row in role_payload[role_key]["stat_summary"].iterrows()
        }
    qualified_teams = con.execute("SELECT COUNT(*) FROM analytics_ti2026_teams").fetchone()[0]
    player_maps = con.execute("SELECT COUNT(*) FROM analytics_player_maps").fetchone()[0]
    banner_profile = fetch_banner_profile(con)
    optimizer_df = fetch_optimizer_recommendations(con)
    con.close()

    write_markdown(
        md_out,
        qualified_teams=qualified_teams,
        player_maps=player_maps,
        role_payload=role_payload,
        banner_profile=banner_profile,
        optimizer_df=optimizer_df,
    )
    write_docx(
        docx_out,
        qualified_teams=qualified_teams,
        player_maps=player_maps,
        role_payload=role_payload,
        banner_profile=banner_profile,
        optimizer_df=optimizer_df,
    )
    print(
        f"Built guide files:\n- {docx_out}\n- {md_out}"
    )


if __name__ == "__main__":
    main()
